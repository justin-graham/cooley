"""
Document Generation Module - Creates Word documents for Minute Book Index and Critical Issues Report.

Generates professional legal documents with exact citations and template-based resolutions.
"""

import io
import logging
from datetime import datetime
from typing import Dict, Any, List
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from app import db

logger = logging.getLogger(__name__)


# Resolution templates for common compliance issues
ISSUE_RESOLUTION_TEMPLATES = {
    'missing_83b': [
        'Consult with tax attorney regarding late 83(b) filing options',
        'Evaluate potential tax liability for affected shareholder',
        'Implement 83(b) verification process for all future grants',
        'Consider automated cap table tool (e.g., Carta) for reminders'
    ],
    'shares_exceed_authorized': [
        'File Amended and Restated Certificate of Incorporation',
        'Obtain board approval for charter amendment',
        'File amendment with state Secretary of State',
        'Retroactively validate excess shares',
        'Authorize buffer shares (10M+) for future issuances'
    ],
    'missing_board_approval': [
        'Verify if approval exists but missing from document set',
        'Obtain unanimous written consent retroactively if needed',
        'Consult with corporate counsel regarding potential liability',
        'Implement approval-required-first process for future issuances'
    ],
    'amendment_approval_threshold': [
        'Contact missing investors to obtain required signatures',
        'Verify whether threshold percentage requirement was met',
        'If threshold not met, consider alternative structures',
        'Consult counsel regarding consequences of deficient approval'
    ],
    'missing_charter': [
        'Locate Certificate of Incorporation from state filing records',
        'Request certified copy from state Secretary of State',
        'Verify company legal name and authorized shares',
        'Update corporate minute book with filed charter document'
    ],
    'default': [
        'Review issue details with corporate counsel',
        'Evaluate legal and compliance risks',
        'Develop remediation plan with specific deadlines',
        'Implement preventive measures for future compliance'
    ]
}


def get_issue_resolution(issue: Dict[str, Any]) -> List[str]:
    """
    Map issue to appropriate resolution template based on category/description.

    Args:
        issue: Issue dict with 'category' and 'description' fields

    Returns:
        List of recommended resolution steps
    """
    category = issue.get('category', '').lower()
    description = issue.get('description', '').lower()

    # Match based on keywords in category or description
    if '83(b)' in description or '83(b)' in category:
        return ISSUE_RESOLUTION_TEMPLATES['missing_83b']
    elif 'exceed' in description and 'authorized' in description:
        return ISSUE_RESOLUTION_TEMPLATES['shares_exceed_authorized']
    elif 'board approval' in description or 'board consent' in description:
        return ISSUE_RESOLUTION_TEMPLATES['missing_board_approval']
    elif 'amendment' in description and ('threshold' in description or 'approval' in description):
        return ISSUE_RESOLUTION_TEMPLATES['amendment_approval_threshold']
    elif 'charter' in category.lower() and 'missing' in description:
        return ISSUE_RESOLUTION_TEMPLATES['missing_charter']
    else:
        return ISSUE_RESOLUTION_TEMPLATES['default']


def format_date_display(date_str: str) -> str:
    """
    Format date string for display (YYYY-MM-DD -> Month D, YYYY).

    Args:
        date_str: Date in YYYY-MM-DD format

    Returns:
        Formatted date string (e.g., "March 15, 2022")
    """
    if not date_str:
        return "Date unknown"

    try:
        date_obj = datetime.strptime(date_str, '%Y-%m-%d')
        return date_obj.strftime('%B %d, %Y')
    except:
        return date_str


def add_paragraph_with_style(doc: Document, text: str, style: str = 'Normal', bold: bool = False,
                              italic: bool = False, font_size: int = 11,
                              alignment: WD_PARAGRAPH_ALIGNMENT = None) -> None:
    """
    Add a paragraph with specific formatting to the document.

    Args:
        doc: Document object
        text: Paragraph text
        style: Paragraph style name
        bold: Bold text
        italic: Italic text
        font_size: Font size in points
        alignment: Paragraph alignment
    """
    para = doc.add_paragraph(text, style=style)
    run = para.runs[0] if para.runs else para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(font_size)
    run.font.name = 'Arial'

    if alignment:
        para.alignment = alignment


def generate_minute_book(audit_id: str) -> bytes:
    """
    Generate Minute Book Index as a Word document.

    Chronological corporate history with exact document citations.

    Args:
        audit_id: UUID of the audit

    Returns:
        Word document as bytes
    """
    try:
        # Fetch audit data from database
        audit = db.get_audit(audit_id)
        if not audit:
            raise ValueError(f"Audit {audit_id} not found")

        company_name = audit.get('company_name', 'Company')
        timeline = audit.get('timeline', [])
        documents = audit.get('documents', [])

        # Create Word document
        doc = Document()

        # Set document margins (1 inch all sides)
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

        # Title page
        title = doc.add_heading('MINUTE BOOK INDEX', level=1)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        for run in title.runs:
            run.font.size = Pt(18)
            run.font.bold = True

        company_heading = doc.add_heading(company_name, level=2)
        company_heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        for run in company_heading.runs:
            run.font.size = Pt(14)

        # Metadata
        doc.add_paragraph(f"Generated: {datetime.now().strftime('%B %d, %Y')}")
        doc.add_paragraph(f"Total Documents: {len(documents)}")
        doc.add_paragraph(f"Total Events: {len(timeline)}")
        doc.add_paragraph("")  # Spacing

        # Divider
        doc.add_paragraph("_" * 80)
        doc.add_paragraph("")

        # Chronological Timeline
        doc.add_heading("CHRONOLOGICAL CORPORATE HISTORY", level=2)
        doc.add_paragraph("")

        # Sort timeline by date
        sorted_timeline = sorted(timeline, key=lambda x: x.get('date', ''))

        for i, event in enumerate(sorted_timeline, start=1):
            date_str = event.get('date', 'Unknown date')
            description = event.get('description', 'No description')
            event_type = event.get('event_type', 'event')
            source_docs = event.get('source_docs', [])

            # Event number and type
            event_label = doc.add_paragraph()
            event_run = event_label.add_run(f"Event #{i} | {event_type.upper().replace('_', ' ')}")
            event_run.font.size = Pt(10)
            event_run.font.bold = True
            event_run.font.color.rgb = RGBColor(212, 43, 30)  # Red accent

            # Date
            date_para = doc.add_paragraph()
            date_run = date_para.add_run(f"Date: {format_date_display(date_str)}")
            date_run.font.size = Pt(11)
            date_run.font.bold = True

            # Description
            desc_para = doc.add_paragraph()
            desc_run = desc_para.add_run(f"Event: {description}")
            desc_run.font.size = Pt(11)

            # Source documents
            if source_docs:
                source_para = doc.add_paragraph()
                source_run = source_para.add_run(f"Source Document: {source_docs[0]}")
                source_run.font.size = Pt(10)
                source_run.italic = True

                # Try to find paragraph reference from extracted data
                # Look for matching event in documents
                para_num = _find_paragraph_reference(audit, event)
                if para_num:
                    para_ref = doc.add_paragraph()
                    para_ref_run = para_ref.add_run(f"Paragraph {para_num}:")
                    para_ref_run.font.name = 'Courier New'
                    para_ref_run.font.size = Pt(9)
                    para_ref_run.font.color.rgb = RGBColor(128, 128, 128)

                    # Add quote if available
                    quote = _find_source_quote(audit, event)
                    if quote:
                        quote_para = doc.add_paragraph()
                        quote_run = quote_para.add_run(f'"{quote}"')
                        quote_run.font.size = Pt(10)
                        quote_run.italic = True
                        quote_para.paragraph_format.left_indent = Inches(0.5)

            # Divider between events
            doc.add_paragraph("─" * 80)
            doc.add_paragraph("")

        # Save to bytes
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.read()

    except Exception as e:
        logger.error(f"Failed to generate minute book: {e}", exc_info=True)
        raise


def generate_issues_report(audit_id: str) -> bytes:
    """
    Generate Critical Issues Report as a Word document.

    Args:
        audit_id: UUID of the audit

    Returns:
        Word document as bytes
    """
    try:
        # Fetch audit data from database
        audit = db.get_audit(audit_id)
        if not audit:
            raise ValueError(f"Audit {audit_id} not found")

        company_name = audit.get('company_name', 'Company')
        issues = audit.get('issues', [])

        # Create Word document
        doc = Document()

        # Set document margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

        # Title page
        title = doc.add_heading('CRITICAL ISSUES REPORT', level=1)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        for run in title.runs:
            run.font.size = Pt(18)
            run.font.bold = True

        company_heading = doc.add_heading(company_name, level=2)
        company_heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        for run in company_heading.runs:
            run.font.size = Pt(14)

        # Metadata
        doc.add_paragraph(f"Generated: {datetime.now().strftime('%B %d, %Y')}")
        doc.add_paragraph(f"Issues Detected: {len(issues)}")
        doc.add_paragraph("")

        # Divider
        doc.add_paragraph("_" * 80)
        doc.add_paragraph("")

        if not issues:
            no_issues_para = doc.add_paragraph()
            no_issues_run = no_issues_para.add_run("✓ No critical issues detected.")
            no_issues_run.font.size = Pt(14)
            no_issues_run.font.bold = True
            no_issues_run.font.color.rgb = RGBColor(0, 128, 0)  # Green
        else:
            # Count by severity
            critical_count = sum(1 for issue in issues if issue.get('severity') == 'critical')
            warning_count = sum(1 for issue in issues if issue.get('severity') == 'warning')
            note_count = sum(1 for issue in issues if issue.get('severity') == 'note')

            summary_para = doc.add_paragraph()
            summary_run = summary_para.add_run(f"Critical: {critical_count} | Warnings: {warning_count} | Notes: {note_count}")
            summary_run.font.size = Pt(11)
            summary_run.font.bold = True

            doc.add_paragraph("")
            doc.add_paragraph("─" * 80)
            doc.add_paragraph("")

            # Each issue
            for i, issue in enumerate(issues, start=1):
                severity = issue.get('severity', 'note').upper()
                category = issue.get('category', 'Unknown')
                description = issue.get('description', 'No description provided')
                source_doc = issue.get('source_doc', '')

                # Issue header
                issue_header = doc.add_paragraph()
                issue_run = issue_header.add_run(f"ISSUE #{i}: {category}")
                issue_run.font.size = Pt(12)
                issue_run.font.bold = True

                # Severity indicator
                severity_para = doc.add_paragraph()
                severity_run = severity_para.add_run(f"Severity: ⚠ {severity}")
                severity_run.font.size = Pt(11)
                if severity == 'CRITICAL':
                    severity_run.font.color.rgb = RGBColor(212, 43, 30)  # Red
                elif severity == 'WARNING':
                    severity_run.font.color.rgb = RGBColor(255, 165, 0)  # Orange
                else:
                    severity_run.font.color.rgb = RGBColor(100, 100, 100)  # Gray

                # Category
                category_para = doc.add_paragraph()
                category_run = category_para.add_run(f"Category: {category}")
                category_run.font.size = Pt(10)

                doc.add_paragraph("")

                # Description
                desc_heading = doc.add_paragraph()
                desc_heading_run = desc_heading.add_run("Description:")
                desc_heading_run.font.bold = True
                desc_heading_run.font.size = Pt(11)

                desc_para = doc.add_paragraph()
                desc_run = desc_para.add_run(description)
                desc_run.font.size = Pt(11)
                desc_para.paragraph_format.left_indent = Inches(0.25)

                doc.add_paragraph("")

                # Recommended Resolution
                resolution_heading = doc.add_paragraph()
                resolution_heading_run = resolution_heading.add_run("Recommended Resolution:")
                resolution_heading_run.font.bold = True
                resolution_heading_run.font.size = Pt(11)

                resolutions = get_issue_resolution(issue)
                for j, resolution_step in enumerate(resolutions, start=1):
                    res_para = doc.add_paragraph()
                    res_run = res_para.add_run(f"{j}. {resolution_step}")
                    res_run.font.size = Pt(10)
                    res_para.paragraph_format.left_indent = Inches(0.25)

                doc.add_paragraph("")

                # Source documents
                if source_doc:
                    source_para = doc.add_paragraph()
                    source_run = source_para.add_run(f"Source Documents:")
                    source_run.font.bold = True
                    source_run.font.size = Pt(10)

                    doc_para = doc.add_paragraph()
                    doc_run = doc_para.add_run(f"• {source_doc}")
                    doc_run.font.size = Pt(10)
                    doc_para.paragraph_format.left_indent = Inches(0.25)

                # Divider
                doc.add_paragraph("")
                doc.add_paragraph("─" * 80)
                doc.add_paragraph("")

            # Summary section
            doc.add_heading("SUMMARY", level=2)
            doc.add_paragraph(f"Critical Issues: {critical_count}")
            doc.add_paragraph(f"Warnings: {warning_count}")
            doc.add_paragraph(f"Notes: {note_count}")
            doc.add_paragraph("")

            next_steps = doc.add_paragraph()
            next_steps_run = next_steps.add_run("Next Steps:")
            next_steps_run.font.bold = True
            next_steps_run.font.size = Pt(11)

            steps = [
                "1. Review with corporate counsel",
                "2. Prioritize critical issues for immediate action",
                "3. Create remediation plan with specific deadlines",
                "4. Implement preventive processes for ongoing compliance"
            ]
            for step in steps:
                step_para = doc.add_paragraph()
                step_run = step_para.add_run(step)
                step_run.font.size = Pt(10)
                step_para.paragraph_format.left_indent = Inches(0.25)

        # Save to bytes
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.read()

    except Exception as e:
        logger.error(f"Failed to generate issues report: {e}", exc_info=True)
        raise


def generate_minute_book_preview(audit_id: str) -> str:
    """
    Generate text preview (first 200 words) of Minute Book Index.

    Args:
        audit_id: UUID of the audit

    Returns:
        Text preview string
    """
    try:
        audit = db.get_audit(audit_id)
        if not audit:
            return "Audit not found."

        company_name = audit.get('company_name', 'Company')
        timeline = audit.get('timeline', [])
        documents = audit.get('documents', [])

        preview = f"""MINUTE BOOK INDEX
{company_name}

Generated: {datetime.now().strftime('%B %d, %Y')}
Total Documents: {len(documents)}
Total Events: {len(timeline)}

CHRONOLOGICAL CORPORATE HISTORY
════════════════════════════════

"""

        # Add first few events
        sorted_timeline = sorted(timeline, key=lambda x: x.get('date', ''))
        for i, event in enumerate(sorted_timeline[:5], start=1):
            date_str = event.get('date', 'Unknown')
            description = event.get('description', 'No description')
            event_type = event.get('event_type', 'event')

            preview += f"Event #{i} | {event_type.upper().replace('_', ' ')}\n"
            preview += f"Date: {format_date_display(date_str)}\n"
            preview += f"Event: {description}\n\n"

        preview += "\n... [Full document contains all events with source citations]"

        # Truncate to ~200 words
        words = preview.split()
        if len(words) > 200:
            preview = ' '.join(words[:200]) + '\n\n... [truncated]'

        return preview

    except Exception as e:
        logger.error(f"Failed to generate minute book preview: {e}")
        return "Preview unavailable"


def generate_issues_preview(audit_id: str) -> str:
    """
    Generate text preview (first 200 words) of Critical Issues Report.

    Args:
        audit_id: UUID of the audit

    Returns:
        Text preview string
    """
    try:
        audit = db.get_audit(audit_id)
        if not audit:
            return "Audit not found."

        company_name = audit.get('company_name', 'Company')
        issues = audit.get('issues', [])

        if not issues:
            return f"""CRITICAL ISSUES REPORT
{company_name}

✓ No critical issues detected.

All documents appear to be in compliance with corporate governance requirements."""

        critical_count = sum(1 for issue in issues if issue.get('severity') == 'critical')
        warning_count = sum(1 for issue in issues if issue.get('severity') == 'warning')

        preview = f"""CRITICAL ISSUES REPORT
{company_name}

Issues Detected: {len(issues)}
Critical: {critical_count} | Warnings: {warning_count}

════════════════════════════════

ISSUE #1: {issues[0].get('category', 'Unknown')}
Severity: {issues[0].get('severity', 'unknown').upper()}

Description:
{issues[0].get('description', 'No description')[:300]}

Recommended Resolution:
"""

        resolutions = get_issue_resolution(issues[0])
        for i, res in enumerate(resolutions[:3], start=1):
            preview += f"{i}. {res}\n"

        preview += "\n... [Full report contains all issues with detailed resolutions]"

        # Truncate to ~200 words
        words = preview.split()
        if len(words) > 200:
            preview = ' '.join(words[:200]) + '\n\n... [truncated]'

        return preview

    except Exception as e:
        logger.error(f"Failed to generate issues preview: {e}")
        return "Preview unavailable"


# Helper functions

def _get_extraction_payload(doc: Dict[str, Any]) -> Dict[str, Any]:
    """
    Support both legacy extracted_data shape and versioned envelope shape.
    """
    extracted_data = doc.get('extracted_data', {})
    if isinstance(extracted_data, dict) and isinstance(extracted_data.get('extraction'), dict):
        return extracted_data.get('extraction', {})
    return extracted_data if isinstance(extracted_data, dict) else {}


def _find_paragraph_reference(audit: Dict[str, Any], event: Dict[str, Any]) -> int:
    """
    Find paragraph number for an event by looking up extracted data.

    Args:
        audit: Full audit data
        event: Timeline event

    Returns:
        Paragraph number or None
    """
    try:
        documents = audit.get('documents', [])
        source_docs = event.get('source_docs', [])
        if not source_docs:
            return None

        source_filename = source_docs[0]

        # Find matching document
        for doc in documents:
            if doc.get('filename') == source_filename:
                extracted_data = _get_extraction_payload(doc)

                # Check for paragraph_number in various extraction types
                if 'stock_issuances' in extracted_data:
                    for issuance in extracted_data['stock_issuances']:
                        if 'paragraph_number' in issuance:
                            return issuance['paragraph_number']

                if 'safe_data' in extracted_data:
                    safe_data = extracted_data['safe_data']
                    if 'paragraph_number' in safe_data:
                        return safe_data['paragraph_number']

                if 'option_data' in extracted_data:
                    option_data = extracted_data['option_data']
                    if 'paragraph_number' in option_data:
                        return option_data['paragraph_number']

                if 'repurchase_data' in extracted_data:
                    repurchase_data = extracted_data['repurchase_data']
                    if 'paragraph_number' in repurchase_data:
                        return repurchase_data['paragraph_number']

        return None

    except Exception as e:
        logger.warning(f"Failed to find paragraph reference: {e}")
        return None


def _find_source_quote(audit: Dict[str, Any], event: Dict[str, Any]) -> str:
    """
    Find verbatim quote for an event by looking up extracted data.

    Args:
        audit: Full audit data
        event: Timeline event

    Returns:
        Source quote or None
    """
    try:
        documents = audit.get('documents', [])
        source_docs = event.get('source_docs', [])
        if not source_docs:
            return None

        source_filename = source_docs[0]

        # Find matching document
        for doc in documents:
            if doc.get('filename') == source_filename:
                extracted_data = _get_extraction_payload(doc)

                # Check for source_quote in various extraction types
                if 'stock_issuances' in extracted_data:
                    for issuance in extracted_data['stock_issuances']:
                        if 'source_quote' in issuance:
                            return issuance['source_quote'][:500]  # Limit to 500 chars

                if 'safe_data' in extracted_data:
                    safe_data = extracted_data['safe_data']
                    if 'source_quote' in safe_data:
                        return safe_data['source_quote'][:500]

                if 'option_data' in extracted_data:
                    option_data = extracted_data['option_data']
                    if 'source_quote' in option_data:
                        return option_data['source_quote'][:500]

                if 'repurchase_data' in extracted_data:
                    repurchase_data = extracted_data['repurchase_data']
                    if 'source_quote' in repurchase_data:
                        return repurchase_data['source_quote'][:500]

        return None

    except Exception as e:
        logger.warning(f"Failed to find source quote: {e}")
        return None
